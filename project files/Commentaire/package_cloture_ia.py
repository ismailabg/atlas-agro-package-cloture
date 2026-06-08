#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=====================================================================
 ATLAS AGRO HOLDING — Commentaires de gestion automatises (IA)
=====================================================================
 Pipeline :
   1. Lecture des 3 extractions ERP (CSV)
   2. Agregation par Business Unit x Poste (reproduit AGREG_PQ)
   3. Detection des ecarts depassant le seuil de materialite
   4. Generation des commentaires de gestion via l'API Claude
      (avec generateur de secours si pas de cle API)
   5. Sorties : rapport PDF, document Word editable, injection Excel

 Usage :
   export ANTHROPIC_API_KEY="sk-ant-..."
   python package_cloture_ia.py --mois 3 --seuil 50

 Auteur : Ismail Abgar
=====================================================================
"""

import os
import sys
import argparse
import shutil
from datetime import datetime

import pandas as pd

# --------------------------------------------------------------------
# PARAMETRES PAR DEFAUT (modifiables en ligne de commande)
# --------------------------------------------------------------------
DOSSIER_CSV    = "./data"                      # dossier des 3 CSV
DOSSIER_SORTIE = "./sorties"                   # dossier des livrables
EXCEL_MASTER   = "./Atlas_Agro_Command_Center.xlsx"  # fichier Excel statique a commenter
MODELE_IA      = "claude-haiku-4-5"            # economique ; "claude-sonnet-4-6" pour + de qualite
DEVISE         = "k€"

MOIS_NOMS = ["Janvier","Fevrier","Mars","Avril","Mai","Juin",
             "Juillet","Aout","Septembre","Octobre","Novembre","Decembre"]

POSTES_EBITDA = ["Chiffre d'affaires","Coût des matières","Charges de personnel",
                 "Charges externes","Transport & logistique","Marketing & publicité","Autres charges"]
POSTES_RN = POSTES_EBITDA + ["Amortissements","Charges financières","Impôt"]


# ====================================================================
# 1. CHARGEMENT DES DONNEES
# ====================================================================
def charger_donnees(dossier):
    """Lit les 3 extractions ERP et renvoie 3 DataFrames."""
    def lire(nom):
        chemin = os.path.join(dossier, nom)
        if not os.path.exists(chemin):
            raise FileNotFoundError(f"Fichier introuvable : {chemin}")
        return pd.read_csv(chemin, sep=";", encoding="utf-8-sig")
    reel26   = lire("GL_Reel_2026.csv")
    budget26 = lire("GL_Budget_2026.csv")
    reel25   = lire("GL_Reel_2025.csv")
    return reel26, budget26, reel25


# ====================================================================
# 2. AGREGATION (reproduit la logique de AGREG_PQ)
# ====================================================================
def agreger(reel26, budget26, reel25, mois):
    """Agrege par Entite x Poste pour le mois de cloture donne."""
    def somme(df, mode):
        d = df[df["Mois"] == mois] if mode == "M" else df[df["Mois"] <= mois]
        return d.groupby(["CodeEntite", "Entite", "PosteP&L"])["Montant_kEUR"].sum()

    base = (budget26[["CodeEntite", "Entite", "PosteP&L"]]
            .drop_duplicates().set_index(["CodeEntite", "Entite", "PosteP&L"]))
    agg = base.copy()
    agg["Reel_M"]      = somme(reel26,   "M")
    agg["Budget_M"]    = somme(budget26, "M")
    agg["Reel_YTD"]    = somme(reel26,   "YTD")
    agg["Budget_YTD"]  = somme(budget26, "YTD")
    agg["Reel_N1_YTD"] = somme(reel25,   "YTD")
    agg = agg.fillna(0).reset_index()
    agg["Ecart_YTD"] = (agg["Reel_YTD"] - agg["Budget_YTD"]).round(1)
    agg["Ecart_M"]   = (agg["Reel_M"]   - agg["Budget_M"]).round(1)
    return agg


def kpis_groupe(agg):
    """Calcule les indicateurs cles consolides (YTD)."""
    def total(poste, col):
        return agg.loc[agg["PosteP&L"] == poste, col].sum()
    ca_r   = total("Chiffre d'affaires", "Reel_YTD")
    ca_b   = total("Chiffre d'affaires", "Budget_YTD")
    ca_n1  = total("Chiffre d'affaires", "Reel_N1_YTD")
    ebitda_r = agg.loc[agg["PosteP&L"].isin(POSTES_EBITDA), "Reel_YTD"].sum()
    ebitda_b = agg.loc[agg["PosteP&L"].isin(POSTES_EBITDA), "Budget_YTD"].sum()
    rn_r = agg.loc[agg["PosteP&L"].isin(POSTES_RN), "Reel_YTD"].sum()
    rn_b = agg.loc[agg["PosteP&L"].isin(POSTES_RN), "Budget_YTD"].sum()
    mb_r = ca_r + total("Coût des matières", "Reel_YTD")
    def pct(a, b): return (a - b) / abs(b) if b else 0
    return {
        "CA_reel": ca_r, "CA_budget": ca_b,
        "CA_vs_budget": pct(ca_r, ca_b), "CA_vs_n1": pct(ca_r, ca_n1),
        "Marge_brute": mb_r, "Marge_brute_pct": mb_r / ca_r if ca_r else 0,
        "EBITDA_reel": ebitda_r, "EBITDA_vs_budget": pct(ebitda_r, ebitda_b),
        "EBITDA_marge": ebitda_r / ca_r if ca_r else 0,
        "RN_reel": rn_r, "RN_vs_budget": pct(rn_r, rn_b),
        "RN_marge": rn_r / ca_r if ca_r else 0,
    }


def ecarts_materiels(agg, seuil):
    """Renvoie les ecarts YTD depassant le seuil, tries par ampleur."""
    mat = agg[abs(agg["Ecart_YTD"]) >= seuil].copy()
    mat["abs"] = mat["Ecart_YTD"].abs()
    mat = mat.sort_values("abs", ascending=False).drop(columns="abs")
    return mat


# ====================================================================
# 3. GENERATION DES COMMENTAIRES (API Claude + secours)
# ====================================================================
def construire_prompt(kpis, mat, mois, annee):
    """Construit le message utilisateur transmis a Claude."""
    lignes = []
    lignes.append(f"Periode de cloture : {MOIS_NOMS[mois-1]} {annee} (cumul YTD).")
    lignes.append("")
    lignes.append("Indicateurs consolides du groupe (YTD) :")
    lignes.append(f"- Chiffre d'affaires : {kpis['CA_reel']:,.0f} {DEVISE} "
                  f"({kpis['CA_vs_budget']:+.1%} vs budget, {kpis['CA_vs_n1']:+.1%} vs N-1)")
    lignes.append(f"- Marge brute : {kpis['Marge_brute']:,.0f} {DEVISE} ({kpis['Marge_brute_pct']:.1%} du CA)")
    lignes.append(f"- EBITDA : {kpis['EBITDA_reel']:,.0f} {DEVISE} "
                  f"({kpis['EBITDA_vs_budget']:+.1%} vs budget, marge {kpis['EBITDA_marge']:.1%})")
    lignes.append(f"- Resultat net : {kpis['RN_reel']:,.0f} {DEVISE} ({kpis['RN_vs_budget']:+.1%} vs budget)")
    lignes.append("")
    lignes.append("Ecarts significatifs vs budget (YTD), par Business Unit et par poste :")
    for _, r in mat.iterrows():
        sens = "favorable" if r["Ecart_YTD"] > 0 else "defavorable"
        lignes.append(f"- {r['Entite']} | {r['PosteP&L']} : ecart {r['Ecart_YTD']:+,.0f} {DEVISE} ({sens})")
    lignes.append("")
    lignes.append("Redige : (1) une SYNTHESE de 2 a 3 phrases pour la direction, "
                  "puis (2) un COMMENTAIRE concis (1 a 2 phrases) par ecart significatif ci-dessus. "
                  "Format attendu :\nSYNTHESE: ...\nCOMMENTAIRES:\n- <BU> — <poste> : <commentaire>")
    return "\n".join(lignes)


SYSTEME_IA = (
    "Tu es controleur de gestion senior chez Atlas Agro Holding, un groupe "
    "agroalimentaire francais (5 business units : cereales/meunerie, huiles, "
    "produits laitiers, sucre/conserves, boissons/jus). Tu rediges les commentaires "
    "de gestion du package de cloture mensuel. Style : factuel, concis, professionnel, "
    "en francais, sans bla-bla ni superlatifs. Chaque commentaire explique l'ecart et "
    "propose une hypothese de cause plausible liee au secteur (saisonnalite, prix des "
    "matieres premieres, effet volume/prix/mix, pression des marques distributeur). "
    "N'invente aucun chiffre : utilise uniquement ceux fournis."
)


def generer_commentaires_ia(kpis, mat, mois, annee, modele):
    """Appelle l'API Claude. Retourne le texte, ou None en cas d'absence de cle/erreur."""
    cle = os.environ.get("ANTHROPIC_API_KEY")
    if not cle:
        return None
    try:
        import anthropic
        client = anthropic.Anthropic()  # lit ANTHROPIC_API_KEY
        message = client.messages.create(
            model=modele,
            max_tokens=1500,
            system=SYSTEME_IA,
            messages=[{"role": "user", "content": construire_prompt(kpis, mat, mois, annee)}],
        )
        return "".join(bloc.text for bloc in message.content if bloc.type == "text").strip()
    except Exception as e:
        print(f"  [!] API Claude indisponible ({e}). Bascule sur le generateur de secours.")
        return None


def generer_commentaires_secours(kpis, mat):
    """Genere des commentaires par regles (sans IA) pour garantir une sortie."""
    out = []
    sens_groupe = "en retrait" if kpis["CA_vs_budget"] < 0 else "en progression"
    out.append("SYNTHESE: Le chiffre d'affaires du groupe ressort "
               f"{sens_groupe} de {abs(kpis['CA_vs_budget']):.1%} par rapport au budget, "
               f"a {kpis['CA_reel']:,.0f} {DEVISE}. La marge EBITDA s'etablit a "
               f"{kpis['EBITDA_marge']:.1%} et le resultat net affiche un ecart de "
               f"{kpis['RN_vs_budget']:+.1%} vs budget.")
    out.append("COMMENTAIRES:")
    for _, r in mat.iterrows():
        sens = "favorable" if r["Ecart_YTD"] > 0 else "defavorable"
        if r["PosteP&L"] == "Chiffre d'affaires":
            cause = "effet volume et saisonnalite" if r["Ecart_YTD"] < 0 else "dynamique commerciale soutenue"
        elif r["PosteP&L"] == "Coût des matières":
            cause = "evolution du prix des matieres premieres"
        elif r["PosteP&L"] == "Charges de personnel":
            cause = "effet de l'interim et des charges sociales"
        else:
            cause = "evolution des couts operationnels"
        out.append(f"- {r['Entite']} — {r['PosteP&L']} : ecart {sens} de "
                   f"{r['Ecart_YTD']:+,.0f} {DEVISE}, lie principalement a {cause}.")
    return "\n".join(out)


# ====================================================================
# 4. SORTIES : PDF / DOCX / EXCEL
# ====================================================================
def exporter_pdf(chemin, kpis, mat, commentaires, mois, annee):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)

    NAVY = colors.HexColor("#1F3864"); BLUE = colors.HexColor("#2E5496")
    styles = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=styles["Title"], textColor=NAVY, fontSize=18)
    sub = ParagraphStyle("sub", parent=styles["Normal"], textColor=NAVY, fontSize=10)
    sec = ParagraphStyle("sec", parent=styles["Heading2"], textColor=BLUE, fontSize=13, spaceBefore=14)
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=15)

    doc = SimpleDocTemplate(chemin, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm,
                            leftMargin=1.8*cm, rightMargin=1.8*cm)
    el = []
    el.append(Paragraph("Atlas Agro Holding", h))
    el.append(Paragraph(f"Commentaires de gestion — Cloture {MOIS_NOMS[mois-1]} {annee} (YTD) — {DEVISE}", sub))
    el.append(Spacer(1, 0.4*cm))

    # Tableau KPI
    el.append(Paragraph("Indicateurs cles du groupe", sec))
    data = [["Indicateur", "Reel", "vs Budget", "vs N-1 / Marge"],
            ["Chiffre d'affaires", f"{kpis['CA_reel']:,.0f}", f"{kpis['CA_vs_budget']:+.1%}", f"{kpis['CA_vs_n1']:+.1%}"],
            ["Marge brute", f"{kpis['Marge_brute']:,.0f}", "", f"{kpis['Marge_brute_pct']:.1%}"],
            ["EBITDA", f"{kpis['EBITDA_reel']:,.0f}", f"{kpis['EBITDA_vs_budget']:+.1%}", f"{kpis['EBITDA_marge']:.1%}"],
            ["Resultat net", f"{kpis['RN_reel']:,.0f}", f"{kpis['RN_vs_budget']:+.1%}", f"{kpis['RN_marge']:.1%}"]]
    t = Table(data, colWidths=[5*cm, 3.5*cm, 3.5*cm, 4*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), NAVY), ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTSIZE", (0,0), (-1,-1), 9),
        ("ALIGN", (1,0), (-1,-1), "RIGHT"), ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F2F2F2")])]))
    el.append(t)

    # Commentaires
    el.append(Paragraph("Commentaires de gestion", sec))
    for ligne in commentaires.split("\n"):
        ligne = ligne.strip()
        if not ligne: continue
        if ligne.upper().startswith("SYNTHESE"):
            el.append(Paragraph("<b>Synthese</b>", body))
            el.append(Paragraph(ligne.split(":",1)[-1].strip(), body))
        elif ligne.upper().startswith("COMMENTAIRES"):
            el.append(Spacer(1, 0.2*cm)); el.append(Paragraph("<b>Detail par ecart</b>", body))
        else:
            el.append(Paragraph(ligne, body))
    el.append(Spacer(1, 0.6*cm))
    el.append(Paragraph(f"<i>Document genere le {datetime.now():%d/%m/%Y a %H:%M}. "
                        f"Commentaires assistes par IA, a valider par le controle de gestion.</i>",
                        ParagraphStyle("foot", parent=body, fontSize=8, textColor=colors.grey)))
    doc.build(el)


def exporter_docx(chemin, kpis, mat, commentaires, mois, annee):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    d = Document()
    titre = d.add_heading("Atlas Agro Holding", level=0)
    titre.runs[0].font.color.rgb = NAVY
    p = d.add_paragraph(f"Commentaires de gestion — Cloture {MOIS_NOMS[mois-1]} {annee} (YTD) — {DEVISE}")
    p.runs[0].italic = True

    d.add_heading("Indicateurs cles du groupe", level=1)
    tab = d.add_table(rows=1, cols=4); tab.style = "Light Grid Accent 1"
    hdr = tab.rows[0].cells
    for i, t in enumerate(["Indicateur","Reel","vs Budget","vs N-1 / Marge"]): hdr[i].text = t
    rows = [["Chiffre d'affaires", f"{kpis['CA_reel']:,.0f}", f"{kpis['CA_vs_budget']:+.1%}", f"{kpis['CA_vs_n1']:+.1%}"],
            ["Marge brute", f"{kpis['Marge_brute']:,.0f}", "-", f"{kpis['Marge_brute_pct']:.1%}"],
            ["EBITDA", f"{kpis['EBITDA_reel']:,.0f}", f"{kpis['EBITDA_vs_budget']:+.1%}", f"{kpis['EBITDA_marge']:.1%}"],
            ["Resultat net", f"{kpis['RN_reel']:,.0f}", f"{kpis['RN_vs_budget']:+.1%}", f"{kpis['RN_marge']:.1%}"]]
    for r in rows:
        cells = tab.add_row().cells
        for i, v in enumerate(r): cells[i].text = v

    d.add_heading("Commentaires de gestion", level=1)
    for ligne in commentaires.split("\n"):
        ligne = ligne.strip()
        if not ligne: continue
        if ligne.upper().startswith("SYNTHESE"):
            sp = d.add_paragraph(); sp.add_run("Synthese : ").bold = True
            sp.add_run(ligne.split(":",1)[-1].strip())
        elif ligne.upper().startswith("COMMENTAIRES"):
            d.add_paragraph().add_run("Detail par ecart :").bold = True
        else:
            d.add_paragraph(ligne, style="List Bullet")

    foot = d.add_paragraph(f"Document genere le {datetime.now():%d/%m/%Y a %H:%M}. "
                           "Commentaires assistes par IA, a valider par le controle de gestion.")
    foot.runs[0].italic = True; foot.runs[0].font.size = Pt(8)
    d.save(chemin)


def injecter_excel(chemin_master, chemin_sortie, commentaires):
    """Injecte les commentaires dans le bloc 5 de SYNTHESE_BU (sur une COPIE)."""
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Alignment
    if not os.path.exists(chemin_master):
        print(f"  [!] Excel master introuvable ({chemin_master}) : injection ignoree.")
        return False
    shutil.copy(chemin_master, chemin_sortie)
    wb = load_workbook(chemin_sortie)
    if "SYNTHESE_BU" not in wb.sheetnames:
        print("  [!] Onglet SYNTHESE_BU absent : injection ignoree."); return False
    ws = wb["SYNTHESE_BU"]
    # Reperer la ligne du bloc 5
    ligne_bloc = None
    for r in range(1, ws.max_row + 1):
        v = ws.cell(row=r, column=2).value
        if isinstance(v, str) and "COMMENTAIRES DE GESTION" in v.upper():
            ligne_bloc = r; break
    if ligne_bloc is None:
        print("  [!] Bloc 5 introuvable : injection ignoree."); return False
    # Ecrire le texte dans la cellule sous le titre
    cible = ws.cell(row=ligne_bloc + 1, column=2)
    cible.value = commentaires
    cible.font = Font(name="Arial", size=9, color="000000")
    cible.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    wb.save(chemin_sortie)
    return True


# ====================================================================
# 5. ORCHESTRATION
# ====================================================================
def main():
    ap = argparse.ArgumentParser(description="Commentaires de gestion IA - Atlas Agro")
    ap.add_argument("--mois", type=int, default=3, help="Mois de cloture (1-12)")
    ap.add_argument("--annee", type=int, default=2026)
    ap.add_argument("--seuil", type=float, default=50, help="Seuil de materialite (k€)")
    ap.add_argument("--csv", default=DOSSIER_CSV)
    ap.add_argument("--sortie", default=DOSSIER_SORTIE)
    ap.add_argument("--excel", default=EXCEL_MASTER)
    ap.add_argument("--modele", default=MODELE_IA)
    args = ap.parse_args()

    os.makedirs(args.sortie, exist_ok=True)
    print(f"=== Cloture {MOIS_NOMS[args.mois-1]} {args.annee} | seuil {args.seuil:.0f} {DEVISE} ===")

    print("[1/5] Lecture des extractions ERP...")
    reel26, budget26, reel25 = charger_donnees(args.csv)

    print("[2/5] Agregation et calcul des indicateurs...")
    agg  = agreger(reel26, budget26, reel25, args.mois)
    kpis = kpis_groupe(agg)
    mat  = ecarts_materiels(agg, args.seuil)
    print(f"      {len(mat)} ecart(s) significatif(s) detecte(s).")

    print("[3/5] Generation des commentaires...")
    txt = generer_commentaires_ia(kpis, mat, args.mois, args.annee, args.modele)
    if txt:
        print(f"      Commentaires generes par l'API Claude ({args.modele}).")
    else:
        txt = generer_commentaires_secours(kpis, mat)
        print("      Commentaires generes par le module de secours (sans IA).")

    suffixe = f"{MOIS_NOMS[args.mois-1]}_{args.annee}"
    print("[4/5] Export PDF et Word...")
    pdf  = os.path.join(args.sortie, f"Commentaires_Cloture_{suffixe}.pdf")
    docx = os.path.join(args.sortie, f"Commentaires_Cloture_{suffixe}.docx")
    exporter_pdf(pdf, kpis, mat, txt, args.mois, args.annee)
    exporter_docx(docx, kpis, mat, txt, args.mois, args.annee)

    print("[5/5] Injection dans l'Excel...")
    xlsx = os.path.join(args.sortie, f"Atlas_Agro_Command_Center_commente_{suffixe}.xlsx")
    ok = injecter_excel(args.excel, xlsx, txt)

    print("\n=== Termine ===")
    print(f"  PDF   : {pdf}")
    print(f"  Word  : {docx}")
    if ok: print(f"  Excel : {xlsx}")


if __name__ == "__main__":
    main()